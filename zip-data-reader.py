import json
import statistics
import docx

from tkinter import *
from tkinter.filedialog import askopenfilename

# secret setting
generate_sample = False

class ZipGradeReporter:
    def __init__(self, master):
        self.import_filename = None
        
        self.master = master
        master.title("ZipGrade Reporter")

        self.message = "Select CSV Report"
        self.label_text = StringVar()
        self.label_text.set(self.message)
        self.label = Label(master, textvariable=self.label_text)

        self.entry = Entry(master)

        self.guess_button = Button(master, text="Select file", command=self.set_import_path)
        self.reset_button = Button(master, text="Generate Report", command=self.generate)

        self.label.grid(row=0, column=0, columnspan=2, sticky=W+E)
        self.entry.grid(row=1, column=0, columnspan=2, sticky=W+E)
        self.guess_button.grid(row=2, column=0)
        self.reset_button.grid(row=2, column=1)

    def set_import_path(self):
        '''
        Gets path to ZipGrade data file.
        '''
        self.import_filename = askopenfilename()
        self.label_text.set(self.import_filename)
        
    def get_export_path(self):
        '''
        Gets path to save report.
        
        Right now it's hard coded as an empty string so the report will be
        saved in the same file as the program is run.
        '''

        return ""

    def get_export_filename(self, records):
        '''
        Gets path to save report. The report file name is simply the quiz name
        and the export date. If no quiz name exists, then the name will default
        to grade_report
        
        ZipGrade date format: May 02 2018 02:14 PM
        '''

        if generate_sample:
            return "sample_report.docx"
        
        months = {"January": "01",
                  "February": "02",
                  "March": "03",
                  "April": "04",
                  "May": "05",
                  "June": "06",
                  "July": "07",
                  "August": "08",
                  "September": "09",
                  "October": "10",
                  "November": "11",
                  "December": "12"}
                  
        r = records[0]
        
        title = r['QuizName'].strip()
        if len(title) == 0:
            title = "grade_report"

        section = r['QuizClass'].strip()
        if len(section) == 0:
            section = ""

        date = r['DataExported'].split(" ")
        yyyy = date[2]
        mm = months[date[0]]
        dd = date[1]

        hh, mm = date[3].split(":")
        period = date[4]

        temp = title + "_" + section + "_" + yyyy + mm + dd
        filename = ""
        underscore = True

        for c in temp:
            if c.isalnum():
                filename += c
                underscore = False
            elif c == "_" and underscore == False:
                filename += c
                underscore = True
            elif underscore == False:
                filename += "_"
                underscore = True
            
        return filename  + ".docx"
    
    def csv_to_json(self, path):
        '''
        Reads ZipGrade csv export file and stores as JSON data
        '''
        
        with open(path, 'r') as f:
            contents = f.read().splitlines()

        fields = contents[0].split(',')

        records = []

        for line in contents[1:]:
            r = {}
            values = line.split(',')
            
            for field, value in zip(fields, values):
                r[field] = value

            records.append(r)

        #pretty = json.dumps(records, indent=4, sort_keys=False)
        #print(pretty)

        sort_by = lambda k: k['LastName'] + " " + k['FirstName']
        records = sorted(records, key=sort_by , reverse=False)
        
        return records

    def get_raw_scores(self, records):
        scores = []

        for r in records:
            scores.append(int(float(r['EarnedPts'])))
            
        return scores
            
    def get_percentages(self, records):
        scores = []

        for r in records:
            s = float(r['PercentCorrect'])
            scores.append(s)
            
        return scores

    def make_meta_data(self, records, document):
        '''
        asdf
        '''

        r = records[0]
        date_created = r['QuizCreated']
        date_exported = r['DataExported']

        p = document.add_paragraph()
        p.add_run("Date Created: ")
        p.add_run(date_created + "\n")
        p.add_run("Date Exported: ")
        p.add_run(date_exported)

    def make_summary_statistics(self, records, document):
        '''
        asdf
        '''
        r = records[0]
        possible_points = records[0]['PossiblePts']
        num_scores = len(records)
        
        scores = self.get_raw_scores(records)
        percentages = self.get_percentages(records)

        document.add_heading('Summary Statistics', 1)

        p = document.add_paragraph()
        p.add_run("Number of Scores: ")
        p.add_run(str(num_scores) + "\n")
        p.add_run("Points Possible: ")
        p.add_run(str(possible_points))

        mean_raw = round(statistics.mean(scores), 2)
        max_raw = max(scores)
        min_raw = min(scores)

        mean_percent = round(statistics.mean(percentages), 2)
        max_percent = max(percentages)
        min_percent = min(percentages)
        
        p = document.add_paragraph()
        p.add_run("Mean (raw/percent): ")
        p.add_run(str(mean_raw) + " / " + str(mean_percent) + "%\n")
        p.add_run("Max (raw/percent): ")
        p.add_run(str(max_raw) + " / " + str(max_percent) + "%\n")
        p.add_run("Min (raw/percent): ")
        p.add_run(str(min_raw) + " / " + str(min_percent) + "%")
        
    def make_difficulty_analysis(self, records, document, hard_threshold, easy_threshold):
        '''
        Generates a list of questions ranked from most to least difficult
        based on the number of students that miss each question.
        '''
        r = records[0]
        num_questions = int(float(r['PossiblePts']))
        
        if 'Key100' in r:
            sheet_size = 100
        elif 'Key50' in r:
            sheet_size = 50
        else:
            sheet_size = 25

        misses = {}

        for r in records:
            for i in range(1, sheet_size + 1):
                correct = r['Key' + str(i)]
                answer = r['Stu' + str(i)]
                
                if len(correct) > 0:
                    if i not in misses:
                        misses[i] = 0
                        
                    if answer != correct:
                        misses[i] += 1

        difficulty = []
        for k, v in misses.items():
            p = round(v / num_questions * 100, 1)
            difficulty.append((k, v, p))

        sort_by = lambda k: k[1]
        difficulty = sorted(difficulty, key=sort_by , reverse=True)

        document.add_heading('Difficulty Analysis', 1)

        paragraph = document.add_paragraph("Most difficult Questions (at least " + str(hard_threshold) + "% missed)\n")
        for d in difficulty:
            if (d[2] >= hard_threshold):
                q = str(d[0])
                n = str(d[1])
                p = str(d[2])
                paragraph.add_run("\tq=" + q + ", n=" + n + ", %=" + p + "\n")

        paragraph = document.add_paragraph("Easiest Questions (no more than " + str(easy_threshold) + "% missed)\n")
        for d in difficulty:
            if (d[2] <= easy_threshold):
                q = str(d[0])
                n = str(d[1])
                p = str(d[2])
                paragraph.add_run("\tq=" + q + ", n=" + n + ", %=" + p + "\n")
                                   
    def make_cover_page(self, records, document):
        '''
        Creates cover page with summary statistics.
        '''
        r = records[0]
        title = r['QuizName']

        document.add_heading('ZipGrade Score Report', 0)
        document.add_heading(title, 1)
        
        self.make_meta_data(records, document)
        self.make_summary_statistics(records, document)
        self.make_difficulty_analysis(records, document, 10, 2)
        
        document.add_page_break()

    def make_score_summary(self, records, document):
        '''
        Creates summary of indidual student scores.
        '''

        document.add_heading('Individual Scores', 1)
        table = document.add_table(rows=1, cols=4)
        table.style = 'Medium Shading 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Raw'
        hdr_cells[2].text = 'Possible'
        hdr_cells[3].text = 'Percent'

        for r in records:
            row_cells = table.add_row().cells
            row_cells[0].text = r['LastName'] + ", " + r['FirstName']
            row_cells[1].text = r['EarnedPts']
            row_cells[2].text = r['PossiblePts']
            row_cells[3].text = r['PercentCorrect'] + "%"
            
        document.add_page_break()

    def make_individual_report(self, record, document):
        '''
        Creates report for a student
        '''
        
        title = record['QuizName']
        section = record['QuizClass']
        last = record['LastName']
        first = record['FirstName']
        zip_id = record['ZipGradeID']
        earned = record['EarnedPts']
        possible = record['PossiblePts']
        percent = record['PercentCorrect']

        if 'Key100' in record:
            sheet_size = 100
        elif 'Key50' in record:
            sheet_size = 50
        else:
            sheet_size = 25

        result = ""
        wrong = ""
        num_wrong = 0
        for i in range(1, sheet_size + 1):
            correct = record['Key' + str(i)]

            if len(correct) > 0:
                answer = record['Stu' + str(i)]

                if len(answer) == 0:
                    answer = "-"
                    
                num_wrong += 1
                wrong += "\t" + str(i) + ". " + answer

                if answer != correct:
                    wrong += " (" + correct + ")"

                    if i < 10:
                        wrong += "\t"
                else:
                    wrong += "\t"
                
                if num_wrong % 5 == 0:
                    wrong += "\n"
            
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_together = True
            
        paragraph.add_run(last + ", " + first + "\n\n").bold = True
        paragraph.add_run("ID: " + zip_id + "\n")
        paragraph.add_run("Test: " + title + "\n")
        paragraph.add_run("Class: " + section + "\n")
        paragraph.add_run("Raw: " + earned + "/" + possible + "\n")
        paragraph.add_run("Percent: " +  percent + "%\n")
        paragraph.add_run("Response Summary: (Correct)\n")
        paragraph.add_run(wrong + "\n")
        paragraph.add_run("\n")
        
    def generate(self):
        '''
        Creates Word Doc with individual score reports.
        '''
        import_path = self.import_filename

        if import_path != None:
            records = self.csv_to_json(import_path)

            document = docx.Document()
            self.make_cover_page(records, document)
            self.make_score_summary(records, document)
            
            for r in records:
                self.make_individual_report(r, document)

            export_path = self.get_export_path()
            export_filename = self.get_export_filename(records)
            save_path = export_path + export_filename
            
            document.save(save_path)

            print("done")
        else:
            print("select a file first")


root = Tk()
my_gui = ZipGradeReporter(root)
root.mainloop()
