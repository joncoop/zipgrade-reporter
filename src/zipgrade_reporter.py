"""
ZipGrade Reporter is a tool that can process the CSV Export data from ZipGrade
and use it to generate reports in a Microsoft Word format. Reports contain
detailed test statistics, score summaries by class, and individual score reports
for distribution to students.
"""

import docx
import json
import os
import statistics
import sys
import tempfile
import urllib.request
import webbrowser

import matplotlib.pyplot as plt; plt.rcdefaults()
import numpy as np
import matplotlib.pyplot as plt

from docx.shared import Inches
from docx.shared import Pt
from tkinter import *
from tkinter.filedialog import askopenfilename

if getattr(sys, 'frozen', False):
    application_path = sys._MEIPASS + '/'
else:
    application_path = os.path.dirname(__file__) + '/'

software_version = 'v0.9-beta.12'
"""str: Version number of this release."""

version_url = 'https://raw.githubusercontent.com/joncoop/zipgrade-reporter/master/src/version.txt'
""" str: URL of version info, used to check if software is up-to-date."""

help_url = "https://joncoop.github.io/zipgrade-reporter/"
"""str: Support website."""


class Scoresheet:
    """
    Quiz data for a single student.

    Scoresheets contain all meta data for a quiz as well as student responses,
    correct answers, and point values for each question.

    Attributes:
        quiz_name (str): Quiz name
        class_name (str): Class name
        first_name (str): Student's first name
        last_name = (str): Student's first name
        zip_id (str): Student's ZipGrade ID
        external_id (str): Unused field in ZipGrade CSV data (Not sure what it's for)
        earned_points (str): Total points earned
        possible_points (str): Total points possible
        percent_correct (str): Percent correct
        date_created (str): Date quiz was created
        date_exported (str): Date quiz data was exported
        key_version (str): Answer key version
        num_questions (str): Number of questions on quiz
        responses = (list) Number of student responses
    """

    def __init__(self, header_row, data_row):
        """
        Constructor for a Scoresheet.

        Args:
            header_row (str): The top row of the ZipGrade CSV export file.
            data_row (str): A single row containing one student's CSV quiz data.
        """
        
        def remove_quotes(s):
            # Gets rid of quotes around non-numeric CSV data values
            # I really should be using the csv module, but I just learned
            # that it existed.
            if len(s) > 0 and s[0] == '"':
                s = s[1:]
            if len(s) > 0 and s[-1] == '"':
                s = s[:-1]

            return s.strip()
        
        delimiter = ","
        fields = header_row.split(delimiter)
        values = data_row.split(delimiter)
        values = [remove_quotes(v) for v in values]
        
        data = {}

        for f, v in zip(fields, values):
            data[f] = v

        self.quiz_name = data['QuizName']
        self.class_name = data['QuizClass']
        self.first_name = data['FirstName']
        self.last_name = data['LastName']
        self.zip_id = data['StudentID']
        self.external_id = data['CustomID'] # unused
        self.earned_points = data['Earned Points']
        self.possible_points = data['Possible Points']
        self.percent_correct = data['PercentCorrect']
        self.date_created = data['QuizCreated']
        self.date_exported = data['DataExported']
        self.key_version = data['Key Version']
        self.num_questions = int((len(fields) - 12) / 4) # 12 metadata colums, 4 data cells per answer

        self.responses = []

        q = 1
        count = 0

        while count < self.num_questions:
            k = 'PriKey' + str(q)
            s = 'Stu' + str(q)

            if k in data:
                student_answer = data[s]
                correct_answer = data[k]
                r = {'question': q, 'answer': student_answer, 'correct': correct_answer}
                self.responses.append(r)
                count += 1

            q += 1

    
class Report:
    """
    Processes multiple ZipGrade scoresheets to create score report.

    A report uses the scoresheets to calculate summary statistics as well
    as to generate the report as an MS Word document.

    Attributes:
        scoresheets (list): List of all scoresheets for a quiz.
    """

    def __init__(self, scoresheets):
        """
        Constructor for a Report.

        Args:
            scoresheets (list): A list of Scoresheets.
        """
        self.scoresheets = scoresheets

        sort_by = lambda k: k.last_name + " " + k.first_name
        self.scoresheets = sorted(self.scoresheets, key=sort_by)

    @property
    def versions(self):
        """list: List of all key versions for a quiz."""

        result = []

        for s in self.scoresheets:
            v = s.key_version

            if v not in result:
                result.append(v)

        result.sort()
        return result

    @property
    def classes(self):
        """list: All classes for a quiz."""
        result = []

        for s in self.scoresheets:
            v = s.class_name

            if v not in result:
                result.append(v)

        result.sort()
        return result

    @property
    def raw_scores(self):
        """list: Raw scores for all students."""
        result = []

        for s in self.scoresheets:
            n = float(s.earned_points)
            result.append(n)

        return result

    @property
    def percentages(self):
        """list: Percentages for all students."""
        result = []

        for s in self.scoresheets:
            n = float(s.percent_correct)
            n = round(n)
            result.append(n)

        return result

    def get_sheets_by_class(self, class_name):
        """
        Gets a list of scoresheets filtered by class.

        Args:
            class_name (str): Name of class to get scoresheets for.

        Returns:
            A filtered list of scoresheets.
        """

        result = []
        for s in self.scoresheets:
            if s.class_name == class_name:
                result.append(s)

        return result

    def get_sheets_by_version(self, key_version):
        """
        Gets a list of scoresheets filtered by key version.

        Args:
            key_version (str): Version to get scoresheets for.

        Returns:
            A filtered list of scoresheets.
        """

        result = []
        for s in self.scoresheets:
            if s.key_version == key_version:
                result.append(s)

        return result

    def quartiles(self, num_list):
        """
        Gets quartiles for a set of values.

        Args:
            num_list (list): List of numbers to calculate quartiles for.
        Returns:
            Lower and upper quartiles for a set of numbers.
        """
        nums = num_list.copy()
        nums.sort()

        mid1 = len(nums) // 2
        mid2 = mid1

        if len(nums) % 2 == 1:
            mid2 += 1

        q1 = round(statistics.median(nums[:mid1]), 2)
        q3 = round(statistics.median(nums[mid2:]), 2)

        return q1, q3

    def add_cover_page(self, document):
        """
        Puts cover page on the report.

        Args:
            document (docx.Document): Document for which content is being added.
        """
        
        sheet_1 = self.scoresheets[0]
        
        document.add_heading('ZipGrade Score Report', 0)
  
        p = document.add_paragraph()
        p.add_run("Quiz Name: ")
        p.add_run(sheet_1.quiz_name + "\n")
        p.add_run("Date Created: ")
        p.add_run(sheet_1.date_created + "\n")
        p.add_run("Date Exported: ")
        p.add_run(sheet_1.date_exported)

        p = document.add_paragraph()
        p.add_run("Classes: " + "\n")
        for class_name in self.classes:
            p.add_run("  - " + class_name + "\n")

    def add_summary_statistics(self, document):
        """
        Generates summary statistics and puts them on document.

        Args:
            document (docx.Document): Document for which content is being added.
        """
        sheet_1 = self.scoresheets[0]
        possible_points = sheet_1.possible_points
        num_scores = len(self.scoresheets)

        mean_raw = round(statistics.mean(self.raw_scores), 2)
        mean_pct = round(statistics.mean(self.percentages), 2)

        median_raw = round(statistics.median(self.percentages), 2)
        median_pct = round(statistics.median(self.percentages), 2)

        st_dev_raw = round(statistics.stdev(self.raw_scores), 2)
        st_dev_pct = round(statistics.stdev(self.percentages), 2)

        min_raw = round(min(self.raw_scores), 2)
        max_raw = round(max(self.raw_scores), 2)
        min_pct = round(min(self.percentages), 2)
        max_pct =  round(max(self.percentages), 2)

        q1_raw, q3_raw = self.quartiles(self.raw_scores)
        q1_pct, q3_pct = self.quartiles(self.percentages)

        document.add_heading('Summary Statistics', 1)

        p = document.add_paragraph()
        p.add_run("Number of Scores: ")
        p.add_run(str(num_scores) + "\n")
        p.add_run("Points Possible: ")
        p.add_run(str(possible_points))

        p = document.add_paragraph()
        p.add_run("Mean (raw/percent): ")
        p.add_run(str(mean_raw) + " / " + str(mean_pct) + "%\n")
        p.add_run("Standard Deviation (raw/percent): ")
        p.add_run(str(st_dev_raw) + " / " + str(st_dev_pct) + "%")

        p = document.add_paragraph()
        p.add_run("Max (raw/percent): ")
        p.add_run(str(max_raw) + " / " + str(max_pct) + "%\n")
        p.add_run("Q3 (raw/percent): ")
        p.add_run(str(q3_raw) + " / " + str(q3_pct) + "%\n")
        p.add_run("Median (raw/percent): ")
        p.add_run(str(median_raw) + " / " + str(median_pct) + "%\n")
        p.add_run("Q1 (raw/percent): ")
        p.add_run(str(q1_raw) + " / " + str(q1_pct) + "%\n")
        p.add_run("Min (raw/percent): ")
        p.add_run(str(min_raw) + " / " + str(min_pct) + "%")

    def add_grade_distribution_graph(self, document):
        """
        Puts bar graph of grade distribution on document.

        Args:
            document (docx.Document): Document for which content is being added.
        """
        ranges = []

        for low in range(0, 100, 5):
            rng = str(low) + '-' + str(low + 4)
            ranges.append(rng)
        ranges.append(100)

        counts = [0] * len(ranges)

        for s in self.scoresheets:
            percent = round(float(s.percent_correct))
            index = min(percent // 5, 20)
            counts[index] += 1

        y_pos = np.arange(len(ranges))

        plt.bar(y_pos, counts, align='center', alpha=0.5)
        plt.xticks(y_pos, ranges, rotation='vertical')
        plt.xlabel('Percent correct', labelpad=12)
        plt.ylabel('Number of students')
        plt.tight_layout()
        
        with tempfile.TemporaryDirectory() as tmpdirname:
            plt.savefig(tmpdirname + '/graph.png')
            document.add_heading('Grade Distribution', 1)
            document.add_picture(tmpdirname + '/graph.png')
        
        
    def add_difficulty_analysis(self, document, sheets, version):
        """
        Generates difficulty analysis and puts it on document.

        Args:
            document (docx.Document): Document for which content is being added.
        """
        document.add_heading('Key version: ' + version, 2)

        misses = {}
        num_questions = sheets[0].num_questions

        for s in sheets:
            for r in s.responses:
                q = int(r['question'])
                a = r['answer']
                c = r['correct']

                if len(c) > 0:
                    if q not in misses:
                        misses[q] = 0

                    if a != c:
                        misses[q] += 1

        difficulty = []
        for k, v in misses.items():
            p = round(v / num_questions * 100, 1)
            difficulty.append((k, v, p))

        sort_by = lambda k: k[1]
        difficulty = sorted(difficulty, key=sort_by , reverse=True)

        if len(difficulty) > 10:
            hard_threshold = difficulty[4][2]
            easy_threshold = difficulty[-3][2]

            paragraph = document.add_paragraph("Most difficult Questions (at least " + str(hard_threshold) + "% missed)\n")
            for d in difficulty:
                if (d[2] >= hard_threshold):
                    q = str(d[0])
                    n = str(d[1])
                    p = str(d[2])
                    paragraph.add_run("\tq=" + q + ", n=" + n + ", %=" + p + "\n")

            paragraph = document.add_paragraph("Easiest Questions (no more than " + str(easy_threshold) + "% missed)\n")
            for d in difficulty:
                if (d[2] <= easy_threshold):
                    q = str(d[0])
                    n = str(d[1])
                    p = str(d[2])
                    paragraph.add_run("\tq=" + q + ", n=" + n + ", %=" + p + "\n")
        else:
            for d in difficulty:
                q = str(d[0])
                n = str(d[1])
                p = str(d[2])
                paragraph.add_run("\tq=" + q + ", n=" + n + ", %=" + p + "\n")

    def add_class_summary(self, document, sheets, summary_title=''):
        """
        Generates class and puts it on document.

        Class summary is an alphabetized list of students with raw scores and
        percentages.

        Args:
            document (docx.Document): Document for which content is being added.
        """
        if summary_title != '':
            document.add_heading('Class scores for ' + summary_title, 1)
        else:
            document.add_heading('Class scores', 1)

        table = document.add_table(rows=1, cols=4)
        table.style = 'Medium Shading 1'
        table.cell(0,0).width = Inches(3.0)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Raw'
        hdr_cells[2].text = 'Possible'
        hdr_cells[3].text = 'Percent'

        for s in sheets:
            row_cells = table.add_row().cells
            row_cells[0].text = s.last_name + ", " + s.first_name
            row_cells[1].text = s.earned_points
            row_cells[2].text = s.possible_points

            rounded_percent = round(float(s.percent_correct))
            row_cells[3].text = str(rounded_percent) + "%"

    def add_individual_report_separator(self, document, class_name):
        """
        Generates separator page to put before individual class reports.

        Args:
            document (docx.Document): Document for which content is being added.
        """
        paragraph = document.add_paragraph()
        paragraph.add_run('\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n')
        heading = document.add_heading('Individual student reports for\n' + class_name, 1)
        heading.alignment = 1

    def add_individual_report(self, document, sheet):
        """
        Generates individual score report for document.

        Individual reports contain student data, scores, and a summary of responses
        along with correct answers.

        Args:
            document (docx.Document): Document for which content is being added.
        """
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_together = True
        
        tab_stops = paragraph.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(0.2))
        tab_stops.add_tab_stop(Inches(0.9))
        tab_stops.add_tab_stop(Inches(1.6))
        tab_stops.add_tab_stop(Inches(2.3))
        tab_stops.add_tab_stop(Inches(3.0))
        tab_stops.add_tab_stop(Inches(3.7))
        tab_stops.add_tab_stop(Inches(4.4))
        tab_stops.add_tab_stop(Inches(5.1))
        tab_stops.add_tab_stop(Inches(5.9))
        tab_stops.add_tab_stop(Inches(6.6))

        name = sheet.last_name + ", " + sheet.first_name
        paragraph.add_run(name + " (ID: " + sheet.zip_id + ")\n").bold = True

        test_name = sheet.quiz_name
        if len(sheet.key_version) > 0:
            test_name += " (Key: " + sheet.key_version + ")"
            
        run = paragraph.add_run("Class: " + sheet.class_name + "\n")
        run.font.size = Pt(9)
        run = paragraph.add_run("Test: " + test_name + "\n")
        run.font.size = Pt(9)
        run = paragraph.add_run("Score: " + sheet.percent_correct + "% " +  
                                "(" + sheet.earned_points + "/" + sheet.possible_points  + ")\n")
        run.font.size = Pt(9)
        
        run = paragraph.add_run("Response Summary: Your Answer (Correct)\n")
        run.font.size = Pt(9)
        
        count = 0
        flagged_questions = []
        
        for r in sheet.responses:
            q = str(r['question'])
            a = str(r['answer'])
            c = str(r['correct'])

            if len(c) > 0:
                item = "\t" + q + ". " + a
                if a != c:
                    item += " (" + c + ")"

                run = paragraph.add_run(item)
                run.font.size = Pt(9)

                count += 1
                if count % 10 == 0:
                   paragraph.add_run('\n')

                if len(c) != len(a):
                    flagged_questions.append(q)

                if len(flagged_questions) == 0:
                    flagged = "None"
                else:
                    flagged = str(flagged_questions)[1: -1]
                    
        return name, flagged
        
    def add_flagged_report_list(self, document, flagged_quizzes):
        """
        Generates class and puts it on document.

        Args:
            document (docx.Document): Document for which content is being added.
            flagged_quizzes (tuple): Class name, student name, and list of questions with flagged responses
        """
        document.add_heading('Flagged Reports', 1)

        if len(flagged_quizzes) > 0:
            paragraph = document.add_paragraph()
            paragraph.add_run("Check that the student responses on flagged questions were scanned correctly. " +
                              "Possible reasons include answers not scanned due light marking, stray marks " +
                              "considered responses due to poor erasing, and marks not read due to glare or " +
                              "poor lighting during scanning. Questions inadvertently left blank by students " +
                              "will also be flagged.")
            paragraph.add_run("\n\n")
            
            paragraph.add_run("From within the ZipGrade app, you can 'Review Papers' and 'Edit Answers' to make " +
                              "corrections. Then redownload the CSV file and generate this report again.")
            paragraph.add_run("\n")
            
            table = document.add_table(rows=1, cols=3)
            table.style = 'Medium Shading 1'

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Class'
            hdr_cells[1].text = 'Name'
            hdr_cells[2].text = 'Flagged Questions'

            for q in flagged_quizzes:
                row_cells = table.add_row().cells
                row_cells[0].text = q[0]
                row_cells[1].text = q[1]
                row_cells[2].text = q[2]
        else:
            paragraph = document.add_paragraph()
            paragraph.add_run("No quizzes have been flagged. It appears that all answers were scanned correctly.")
            paragraph.add_run("\n")

    def generate(self):
        """
        Creates a ZipGrade report as a Word document.

        The report contains a cover page with basic quiz information and
        summary statistics. Subsiquent pages include difficlty analysis, class
        summaries, and individual score reports.

        Returns:
            The completed report.
        """
        document = docx.Document()

        # styling
        style = document.styles['Normal']
        font = style.font
        font.size = Pt(11)

        sections = document.sections
        for section in sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
        
        # cover page
        self.add_cover_page(document)
        document.add_page_break()

        # summary statistics
        self.add_summary_statistics(document)
        self.add_grade_distribution_graph(document)
        document.add_page_break()

        # difficulty analysis
        document.add_heading('Difficulty Analysis', 1)
        for version in self.versions:
            sheets = self.get_sheets_by_version(version)
            self.add_difficulty_analysis(document, sheets, version)
        document.add_page_break()

        # class reports
        for class_name in self.classes:
            sheets = self.get_sheets_by_class(class_name)
            self.add_class_summary(document, sheets, class_name)
            document.add_page_break()

        # individual reports
        flagged_quizzes = []
        
        for i, class_name in enumerate(self.classes):
            sheets = self.get_sheets_by_class(class_name)
            self.add_individual_report_separator(document, class_name)
            document.add_page_break()

            for s in sheets:
                name, flags = self.add_individual_report(document, s)

                if flags != "None":
                    flagged_quizzes.append([class_name, name, flags])

            if i + 1 < len(self.classes):
                document.add_page_break()

        document.add_page_break()

        # flagged reports
        self.add_flagged_report_list(document, flagged_quizzes)

        # all done
        return document
         
class App:
    """
    GUI component of ZipGrade Reporter.

    Attributes:
        import_path (str): Path to CSV file.
        export_path (str): Path to save final report.
    """

    def __init__(self, master):
        """Constructor for an App
        """

        self.import_path = None
        self.export_path = None

        self.master = master
        self.gui_init()

    def gui_init(self):
        """
        Defines App layout
        """
        self.master.iconbitmap(application_path + 'images/icon.ico')
        self.master.title("ZipGrade Reporter")
        self.master.resizable(False, False)

        select_button = Button(self.master, text="1. Select ZipGrade CSV Data", command=self.select_file)
        select_button.config(width=30)
        select_button.grid(row=0, column=0, padx=5, pady=5, sticky=(W))

        generate_button = Button(self.master, text="2. Generate Report", command=self.generate)
        generate_button.config(width=30)
        generate_button.grid(row=0, column=1, padx=5, pady=5, sticky=(E))

        instr1 = Label(self.master, text="The following data file will be used to generate your report...")
        instr1.grid(row=3, column=0, columnspan=2, padx=5, pady=5, sticky=(W))

        self.import_lbl_text = StringVar()
        self.import_lbl_text.set("Waiting for file selection...")
        import_lbl = Label(self.master, textvariable=self.import_lbl_text)
        import_lbl.grid(row=4, column=0, columnspan=2, padx=20, pady=5, sticky=(W))

        instr2 = Label(self.master, text="Report will be created in...")
        instr2.grid(row=5, column=0, columnspan=2, padx=5, pady=5, sticky=(W))

        self.export_lbl_text = StringVar()
        self.export_lbl_text.set("...")
        export_lbl = Label(self.master, textvariable=self.export_lbl_text)
        export_lbl.grid(row=6, column=0, columnspan=2, padx=20, pady=5, sticky=(W))

        self.status_lbl_text = StringVar()
        status_lbl = Label(self.master, textvariable=self.status_lbl_text)
        status_lbl.grid(row=7, column=0, columnspan=2, padx=5, pady=5, sticky=(W))

        links = Frame(self.master)

        help_link = Label(links, text="Help", fg="blue", cursor="hand2")
        help_link.pack( side = LEFT )
        help_link.bind("<Button-1>", lambda e: webbrowser.open_new(help_url))

        if not self.is_up_to_date():
            slash = Label(links, text=" | ", fg="gray", cursor="hand2")
            slash.pack( side = LEFT )
                    
            update_link = Label(links, text="Update ZipGrade Reporter", fg="blue", cursor="hand2")
            update_link.pack( side = LEFT )
            update_link.bind("<Button-1>", lambda e: webbrowser.open_new(help_url))

        links.grid(row=9, column=0, columnspan=1, padx=5, pady=5, sticky=(W))

        version = Label(self.master, text=software_version, fg="gray")
        version.grid(row=9, column=1, columnspan=1, padx=5, pady=5, sticky=(E))

    def is_up_to_date(self):
        """
        Checks the ZipGradeReporter website to see if application is latest version.

        Returns:
            True if up-to-date, False otherwise
        """

        try:
            fp = urllib.request.urlopen(version_url)
            mybytes = fp.read()
            version_txt = mybytes.decode('utf8')
            fp.close()

            start_del = "StringStruct(u'FileVersion', u'"
            end_del = "'),"

            start_loc = version_txt.find(start_del) + len(start_del)
            end_loc = version_txt.find(end_del, start_loc)

            version = 'v' + version_txt[start_loc: end_loc]

            if version == software_version:
                return True
        except:
            pass
        
        return False
    
    def select_file(self):
        """
        Sets path to ZipGrade data file and sets export path to same directory.
        """

        self.import_path = askopenfilename()
        self.export_path = os.path.dirname(self.import_path)

        self.import_lbl_text.set(self.import_path)
        self.export_lbl_text.set(self.export_path)

    def change_export_path(self):
        """
        Sets save path for ZipGrade report.

        This feautre is currently unimplemented!
        """
        pass

    def get_export_filename(self, sheet):
        """
        Gets path to save report.

        The report file name is simply the quiz name and the export date. If no
        quiz name exists, then the name will default to grade_report

        Note:
            ZipGrade date format: May 02 2018 02:14 PM (phone)
                                  2019-09-18 00:00:00       (web)


        Attributes:
            sheet (Scoresheet): Single scoresheet to extract quiz data from.
        """

        months = {"Jan": "01", "Feb": "02", "Mar": "03", "Apr": "04",
                  "May": "05", "Jun": "06", "Jul": "07", "Aug": "08",
                  "Sep": "09", "Oct": "10", "Nov": "11", "Dec": "12"}
        
        title = sheet.quiz_name.strip()
        if len(title) == 0:
            title = "ZipGradeReport"

        if "-" in sheet.date_exported:
            date = sheet.date_exported.split("-")
            yyyy = date[0]
            mm = date[1]
            dd = date[2][:2]
        elif "/" in sheet.date_exported:
            date = sheet.date_exported.split("/")
            yyyy = date[2]
            mm = date[0]
            dd = date[1]
        else:
            date = sheet.date_exported.split(" ")
            print(date)
            yyyy = date[2]
            mm = months[date[0]]
            dd = date[1]

        if int(dd) < 10:
            dd = "0" + dd
            
        if int(mm) < 10:
            mm = "0" + mm

        temp = title + "_" + "_" + yyyy + mm + dd
        filename = ""
        underscore = True

        for c in temp:
            if c.isalnum():
                filename += c
                underscore = False
            elif c == "_" and underscore == False:
                filename += c
                underscore = True
            elif underscore == False:
                filename += "_"
                underscore = True

        return filename  + ".docx"

    def save(self, document):
        """
        Sets save path for ZipGrade report.

        Attributes:
            document (docx.Document): Finalized document to save.
        """

        try:
            document.save(self.save_path)
            self.status_lbl_text.set("Your report is ready!")
        except:
            self.status_lbl_text.set("Unable to save report. Check file and disk permissions.")

    def fix_csv(self, header_str):
        """
        Replaces mobile app CSV headers with those from CSV file downloaded from
        ZipGrade website.
        """

        if 'ZipGradeID' in header_str:
            header_str = header_str.replace('ZipGradeID', 'StudentID')
            header_str = header_str.replace('ExternalID', 'CustomID')
            header_str = header_str.replace('EarnedPts', 'Earned Points')
            header_str = header_str.replace('PossiblePts', 'Possible Points')
            header_str = header_str.replace('Key', 'PriKey')
            header_str = header_str.replace('PriKeyVersion', 'Key Version')
            header_str = header_str.replace('EarnedPts', 'Earned Points')
            header_str = header_str.replace('PossPt', 'Mark')

        return header_str
                    
    def generate(self):
        """
        Reads ZipGrade CSV file and generates report.

        Valid CSV files begin with a single line with all data fields. Each subsiquent
        line contains individual student quiz data.
        """
        
        generated = False

        if self.import_path != None:
            try:
                with open(self.import_path) as f:
                    lines = f.readlines()

                header = lines[0]
                header = self.fix_csv(header)

                all_sheets = []

                for line in lines[1:]:
                    sheet = Scoresheet(header, line)
                    all_sheets.append(sheet)

                r = Report(all_sheets)
                document = r.generate()
                generated = True

            except Exception as inst:
                self.status_lbl_text.set("Something went wrong. Be sure your CSV data file is valid.")

            if generated:
                self.save_path = self.export_path + '/' + self.get_export_filename(all_sheets[0])
                self.save(document)
        else:
            self.status_lbl_text.set("You must select a file first!")


# Let's do this!
if __name__ == "__main__":
    root = Tk()
    my_gui = App(root)
    root.mainloop()
